from pathlib import Path
from docx import Document
import re

TEMPLATE = Path(r"D:\ГО Талан UA\Гранти\Конкурс IREX\1. RFA_UVRR_Reginal_Coordination_Project Proposal_template.docx")
# find latest .bak for this template
bak_files = sorted(TEMPLATE.parent.glob(TEMPLATE.name + '.bak.*'))
source = bak_files[-1] if bak_files else TEMPLATE
OUTPUT = TEMPLATE.with_name(TEMPLATE.stem + '_filled.docx')

print('Using source for patching:', source)
print('Will save to:', OUTPUT)

EXEC_SUM_FIRST = (
    'Створення інформаційно-координаційного ХАБу "Новий Шлях" допоможе ветерану не витрачати критично дефіцитні сили, '
    'час, емоції та здоров\'я на хаотичний пошук фахівців, записи на візити та оформлення документів, яких у нього (фізично та ментально) часто просто немає, '
    'особливо якщо є проблеми з пересуванням через ампутації тощо.'
)
ANALYTICAL_BRIEFS = (
    'Показник для Objective 4: Провести 5 аналітичних брифів — по одному звіту кожні 2 місяці після запуску для оперативного донесення виявлених прогалин до органів влади.'
)
BUDGET_NOTE = (
    'Власний нефінансовий внесок ГО «ТАЛАН ЮА» (In-kind contribution): $3,500. '
    'Запитуване фінансування від IREX: $20,000 (виділено в IT-статті бюджету для розгортання, безпеки та тестування).'
)

if not source.exists():
    print('Source not found:', source)
    raise SystemExit(1)

doc = Document(source)

def find_paragraph_index(keys):
    for i,p in enumerate(doc.paragraphs):
        for k in keys:
            if k in p.text:
                return i
    return None

# summary
summary_keys = ['1. Summary', 'Короткий опис', '1. Summary (up to 400 words)']
idx = find_paragraph_index(summary_keys)
if idx is not None:
    j = idx + 1
    while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() == '':
        j += 1
    if j < len(doc.paragraphs):
        para = doc.paragraphs[j]
        if EXEC_SUM_FIRST not in para.text:
            para.text = EXEC_SUM_FIRST + '\n\n' + para.text
            print('Inserted executive summary first sentence at paragraph', j+1)
    else:
        doc.add_paragraph(EXEC_SUM_FIRST)
        print('Appended executive summary at end')
else:
    doc.add_paragraph(EXEC_SUM_FIRST)
    print('Summary heading not found; appended executive summary at end')

# objective 4
obj_keys = ['Objective 4', '4.', 'Objective 4:']
obj_idx = find_paragraph_index(obj_keys)
if obj_idx is not None:
    if 'аналітич' not in doc.paragraphs[obj_idx].text:
        doc.paragraphs[obj_idx].text = doc.paragraphs[obj_idx].text + ' ' + ANALYTICAL_BRIEFS
        print('Updated Objective 4 at paragraph', obj_idx+1)
else:
    doc.add_paragraph(ANALYTICAL_BRIEFS)
    print('Objective 4 not found; appended analytics brief at end')

# budget note append
doc.add_paragraph(BUDGET_NOTE)
print('Appended budget note at end as fallback')

# save
doc.save(OUTPUT)
print('Saved filled output to', OUTPUT)
