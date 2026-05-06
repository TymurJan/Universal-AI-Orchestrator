import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse  # Додано для підтримки командного рядка
import sys

# Шлях до еталонної шапки ГО «ТАЛАН ЮА» (Першоджерело з логотипом та актуальними реквізитами)
HEADER_TEMPLATE_PATH = r"D:\ГО Талан UA\Листи\Шаблони\Шапка.docx"

def add_branding_header(doc, page_width):
    """
    Фірмовий бланк ГО «ТАЛАН ЮА» (чітко 3 рядки, без зайвих переносів).
    """
    if not os.path.exists(HEADER_TEMPLATE_PATH):
        print(f'⚠️ Шапку не знайдено: {HEADER_TEMPLATE_PATH}')
        return

    from io import BytesIO
    src = Document(HEADER_TEMPLATE_PATH)

    # Витягуємо blob логотипу
    logo_blob = None
    for rel in src.part.rels.values():
        if 'image' in rel.reltype:
            logo_blob = rel.target_part.blob
            break

    # Таблиця: 2 колонки
    tbl = doc.add_table(rows=1, cols=2)
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.width = page_width
    tbl.autofit = False
    
    logo_col_w = Inches(1.1)
    text_col_w = page_width - logo_col_w
    
    tbl.columns[0].width = logo_col_w
    tbl.columns[1].width = text_col_w
    for row in tbl.rows:
        row.cells[0].width = logo_col_w
        row.cells[1].width = text_col_w

    # --- Ліва колонка: лого ---
    logo_cell = tbl.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_blob:
        logo_para.add_run().add_picture(BytesIO(logo_blob), width=Inches(0.95))
    logo_para.paragraph_format.space_after = Pt(16)

    # --- Права колонка: 3 ідеальні рядки ---
    text_cell = tbl.cell(0, 1)
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    lines_cfg = [
        ("ГРОМАДСЬКА ОРГАНІЗАЦІЯ «ТАЛАН ЮА»", False, 14),
        ("Р/Р UA493052990000026001031608581 АТ КБ «ПРИВАТБАНК» • ЄДРПОУ 45119390", False, 8),
        ("вул. Волкова, буд. 95, кв. 26, м. Черкаси, 18005, тел.: +380636568999, email: ngo.talan.ua@gmail.com", False, 8)
    ]
    
    for i, (txt, bold, size) in enumerate(lines_cfg):
        p = text_cell.paragraphs[0] if i == 0 else text_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = bold
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        # Додаємо "повітря" (відступи) між рядками за допомогою space_after
        if i == 0:
            p.paragraph_format.space_after = Pt(8)  # Відступ під Назвою ГО
        elif i == 1:
            p.paragraph_format.space_after = Pt(8)  # Відступ під Реквізитами
        else:
            p.paragraph_format.space_after = Pt(16) # Більший відступ до нижньої лінії

    # Приховуємо рамки таблиці (крім нижньої)
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for b_type in ['top', 'left', 'right', 'insideH', 'insideV']:
        edge = OxmlElement(f'w:{b_type}')
        edge.set(qn('w:val'), 'none')
        tblBorders.append(edge)
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), '000000')
    tblBorders.append(bot)
    tblPr.append(tblBorders)

    doc.add_paragraph('').paragraph_format.space_after = Pt(6)









def clean_markdown_all_symbols(text):
    """Видаляє абсолютно всі Markdown символи та застосовує юридичні правки назв"""
    # 0. Попередня чистка від технічного сміття Markdown
    text = re.sub(r'^-+$', '', text) # Видаляємо лінії ---
    text = re.sub(r'\.{3,}', '', text) # Видаляємо довгі точки
    
    # 1. Юридична корекція назв ТГ та ДФ
    text = text.replace('територіальної громади', 'Територіальної Громади')
    text = text.replace('територіальна громада', 'Територіальна Громада')
    text = text.replace('ДФ Черкаської територіальної громади', 'ДФ Черкаської ТГ')
    text = text.replace('ДФТГ', 'ДФ Черкаської ТГ')
    
    # 2. Інтелектуальна заміна Сторiн з урахуванням відмінків (Regex рушій)
    status_map = {
        r'\bОтримувачем\b': 'Стороною 1',
        r'\bОтримувачу\b': 'Стороні 1',
        r'\bОтримувача\b': 'Сторони 1',
        r'\bОтримувач\b': 'Сторона 1',
        r'\bНабувачем\b': 'Стороною 2',
        r'\bНабувачу\b': 'Стороні 2',
        r'\bНабувача\b': 'Сторони 2',
        r'\bНабувач\b': 'Сторона 2'
    }
    for pattern, replacement in status_map.items():
        text = re.sub(pattern, replacement, text)
    
    # 3. Видаляємо стандартні Markdown символи
    text = re.sub(r'^#+\s*', '', text) # Заголовки
    text = re.sub(r'(\*\*|\*|__|_)','', text) # Жирний/курсив
    text = re.sub(r'^\d+\.\s*', '', text) # Нумерація списків
    
    return text.strip()

def md_to_docx(md_path, docx_path):
    # --- [1] ОЦІНКА ОБ'ЄМУ ТЕКСТУ ---
    with open(md_path, 'r', encoding='utf-8') as f:
        full_content = f.read()
        char_count = len(full_content)
    
    # Визначаємо ТИП документа за першим рядком
    first_lines = full_content[:200].upper()
    is_order = 'НАКАЗ' in first_lines or 'ЛИСТ' in first_lines or 'ЗВЕРНЕННЯ' in first_lines
    # Акти приймання-передачі, внутрішні документи — без логотипу
    
    # Стиль верстки
    is_compact = char_count < 3800
    
    doc = Document()
    
    # СТАНДАРТИ ДСТУ 4163:2020 (Поля 30/10/20/20)
    section = doc.sections[0]
    section.top_margin = Inches(0.79) 
    section.bottom_margin = Inches(0.79) 
    section.left_margin = Inches(1.18)   
    section.right_margin = Inches(0.39)  
    
    page_width = section.page_width - section.left_margin - section.right_margin - Pt(5)
    
    # ГЛОБАЛЬНИЙ СТИЛЬ
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    if is_compact:
        style.paragraph_format.line_spacing = 1.15
        p_space_after = Pt(6)
        p_space_before = Pt(0)
    else:
        style.paragraph_format.line_spacing = 1.1
        p_space_after = Pt(8)
        p_space_before = Pt(4)
    
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # ФІРМОВИЙ БЛАНК (тільки для Наказів, Листів, Звернень)
    if is_order:
        add_branding_header(doc, page_width)
    
    lines = full_content.splitlines()
    body_lines = [l.strip() for l in lines if l.strip()]
    
    # Знаходимо останній абзац перед підписами
    last_body_idx = -1
    for i, line in enumerate(body_lines):
        if 'Підписи сторін' in line or 'ПІДПИСИ СТОРІН' in line:
            last_body_idx = i - 1
            break
            
    for i, line in enumerate(body_lines):
        clean_text = clean_markdown_all_symbols(line)
        if not clean_text:
            continue
            
        # [2] ЛОГІКА ВЕРСТКИ (З ПРИОРИТЕТОМ ПІДПИСІВ)
        
        # 1. Блок Підписів Сторiн
        if 'Підписи сторін' in clean_text or 'ПІДПИСИ СТОРІН' in clean_text:
            p_header = doc.add_paragraph('Підписи сторін:')
            p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_header.paragraph_format.space_before = Pt(24)
            p_header.paragraph_format.keep_with_next = True # Не відривати від таблиці
            
            p_space = doc.add_paragraph('')
            p_space.paragraph_format.space_after = Pt(6)
            
            table = doc.add_table(rows=4, cols=2)
            table.width = page_width
            table.autofit = False
            col_w = int(page_width / 2)
            for row in table.rows:
                for cell in row.cells:
                    cell.width = col_w
            
            # Рядок 0: Сторони
            table.cell(0, 0).text = "Сторона 1:"
            table.cell(0, 0).paragraphs[0].runs[0].bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            
            table.cell(0, 1).text = "Сторона 2:"
            table.cell(0, 1).paragraphs[0].runs[0].bold = True
            table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
            
            # Рядок 1: Посади (Тут довжина тексту може відрізнятись, але рядок таблиці їх вирівняє)
            table.cell(1, 0).text = "Голова Правління ГО «ТАЛАН ЮА»"
            table.cell(1, 1).text = "Командир ДФ Черкаської Територіальної Громади №1"
            
            # Рядок 2: Лінії підписів
            p_sig1 = table.cell(2, 0).paragraphs[0]
            p_sig1.text = "________________  /Т.Ю. ШАПОВАЛ/"
            p_sig1.paragraph_format.space_before = Pt(24) # Простір для самого підпису ручкою
            
            p_sig2 = table.cell(2, 1).paragraphs[0]
            p_sig2.text = "________________  /В.А. ГОРА/"
            p_sig2.paragraph_format.space_before = Pt(24)
            
            # Рядок 3: М.П.
            table.cell(3, 0).text = "М.П."
            table.cell(3, 1).text = "М.П. (за наявності)"
            
            for row in table.rows:
                for p in row.cells[0].paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for p in row.cells[1].paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tbl = table._tbl
            tblPr = tbl.xpath('w:tblPr')[0]
            tblBorders = OxmlElement('w:tblBorders')
            for b_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                edge = OxmlElement(f'w:{b_type}')
                edge.set(qn('w:val'), 'none')
                tblBorders.append(edge)
            tblPr.append(tblBorders)
            break 

        # 2. Заголовок (Акт/Наказ)
        elif line.startswith('# '):
            if is_order and 'ГРОМАДСЬКА ОРГАНІЗАЦІЯ' in clean_text.upper():
                continue  # Пропускаємо дублюючу назву ГО, бо вона вже в шапці
            p = doc.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            
        # 3. Шапка (Місце і Дата)
        elif 'м. Черкаси' in clean_text:
            table = doc.add_table(rows=1, cols=2)
            table.width = page_width
            table.autofit = False
            
            date_str = "________________"
            if '«' in line:
                match = re.search(r'«(.*?)»\s*(.*?)\s*р\.', line)
                if match:
                    date_str = f"«{match.group(1)}» {match.group(2)} р."
                
            table.cell(0, 0).text = "м. Черкаси"
            p_date = table.cell(0, 1).paragraphs[0]
            p_date.text = date_str
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tbl = table._tbl
            tblPr = tbl.xpath('w:tblPr')[0]
            tblBorders = OxmlElement('w:tblBorders')
            for b_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                edge = OxmlElement(f'w:{b_type}')
                edge.set(qn('w:val'), 'none')
                tblBorders.append(edge)
            tblPr.append(tblBorders)
            doc.add_paragraph('').paragraph_format.space_after = Pt(12)
            
        # 4. Підзаголовки і звичайний текст
        elif line.startswith('## '):
            p = doc.add_paragraph(clean_text)
            p.runs[0].bold = True
            p.paragraph_format.space_before = p_space_before
            p.paragraph_format.space_after = p_space_after
        else:
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.space_before = p_space_before
            p.paragraph_format.space_after = p_space_after
            
            # ЗАХИСТ: Якщо це останній абзац перед підписами — не відриваємо його
            if i == last_body_idx:
                p.paragraph_format.keep_with_next = True

    doc.save(docx_path)
    print(f"✅ Документ успішно збережено: {docx_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Генератор DOCX з Markdown для ГО ТАЛАН ЮА')
    parser.add_argument('--input', type=str, help='Шлях до вхідного Markdown файлу')
    parser.add_argument('--output', type=str, help='Шлях для збереження DOCX')
    
    args = parser.parse_args()
    
    if args.input and args.output:
        # Автоматичний запуск через аргументи
        md_to_docx(args.input, args.output)
    else:
        # Ручний режим (для відладки)
        print("🚩 Запуск у ручному режимі (Default). Використовуйте --input та --output для автоматизації.")
        # Тут можна залишити дефолтні шляхи для тестів, якщо треба
    print(f"✅ Smart Success: {docx_path} (Style: {'Compact' if is_compact else 'Balanced'})")

if __name__ == "__main__":
    dropzone = r"d:\ГО Талан UA\Talan UA Antigravity manager\📥_DROPZONE"
    
    import glob
    md_files = glob.glob(os.path.join(dropzone, "*.md"))
    
    for md_p in md_files:
        filename = os.path.basename(md_p)
        # Пропускаємо службові та тестові файли
        if any(x in filename for x in ["README", "LONG_TEST", "_test"]):
            continue
        docx_p = md_p.replace(".md", ".docx")
        try:
            md_to_docx(md_p, docx_p)
        except PermissionError:
            print(f"⚠️  Файл зайнятий (відкритий у Word): {filename} — пропускаємо")
        except Exception as e:
            print(f"❌ Помилка в {filename}: {e}")
