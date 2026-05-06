import os
from docx import Document
import shutil

def generate_document(template_path, output_path, content_blocks):
    """
    Creates a new document by copying the template (header/logo/styles)
    and appending new content.
    """
    # 1. Створюємо копію шаблону, щоб зберегти всі стилі, шапку та логотип
    shutil.copy2(template_path, output_path)
    
    # 2. Відкриваємо копію для редагування
    doc = Document(output_path)
    
    # 3. Додаємо контент (текст документа)
    # Ми додаємо порожній рядок для відступу від шапки
    doc.add_paragraph("\n")
    
    for block in content_blocks:
        if block.get('type') == 'paragraph':
            doc.add_paragraph(block.get('text', ''))
        elif block.get('type') == 'heading':
            level = block.get('level', 1)
            style_name = f'Heading {level}'
            try:
                doc.add_heading(block.get('text', ''), level=level)
            except KeyError:
                # Якщо стиль за замовчуванням відсутній, створюємо жирний параграф
                p = doc.add_paragraph()
                run = p.add_run(block.get('text', ''))
                run.bold = True
                # Можна також налаштувати розмір шрифту тут, якщо потрібно
            
    # 4. Зберігаємо результат
    doc.save(output_path)
    print(f"Документ успішно створено: {output_path}")

if __name__ == "__main__":
    # Тестовий запуск
    T_PATH = r"d:\ГО Талан UA\Листи\Шаблони\шапка.docx"
    O_PATH = r"d:\ГО Талан UA\Листи\test_document.docx"
    
    TEST_CONTENT = [
        {'type': 'heading', 'text': 'ТЕСТОВИЙ ДОКУМЕНТ', 'level': 1},
        {'type': 'paragraph', 'text': 'Це автоматично згенерований текст для перевірки збереження шапки та логотипу.'},
        {'type': 'paragraph', 'text': 'Дата: 04 березня 2026 року.'}
    ]
    
    if os.path.exists(T_PATH):
        generate_document(T_PATH, O_PATH, TEST_CONTENT)
    else:
        print(f"Помилка: Шаблон не знайдено за шляхом {T_PATH}")
