from pathlib import Path
from docx import Document

paths = [
    Path(r'D:\ГО Талан UA\Гранти\Конкурс IREX\Робоча\01_IREX_Grant_Proposal_NovyShlyakh_FULL_v6.docx'),
    Path(r'D:\ГО Талан UA\Гранти\Конкурс IREX\Робоча\Заявка\01_IREX_Grant_Proposal_NovyShlyakh_FULL_v7.docx'),
]
for p in paths:
    print('FILE:', p)
    if not p.exists():
        print('  missing')
        continue
    doc = Document(p)
    print('  paragraphs:', len(doc.paragraphs))
    for i, para in enumerate(doc.paragraphs[:40], 1):
        text = para.text.strip()
        if text:
            print(f'{i}: {text}')
    print('---')
