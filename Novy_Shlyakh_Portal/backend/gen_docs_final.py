import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define the data with the new Ukrainian plate
data = {
    "org_name": "ГРОМАДСЬКА ОРГАНІЗАЦІЯ «ТАЛАН ЮА»",
    "edrpou": "45119390",
    "address": "18005, Черкаська обл., м. Черкаси, вул. Волкова, буд. 95",
    "car_brand": "SKODA FABIA",
    "vin": "TMBGC26Y364517308",
    "ua_plate": "CA8063KE",
    "de_plate": "GD MK826",
    "year": "2005",
    "unique_code": "23844937",
    "decl_date": "11.05.2024",
    "reg_date": "13.06.2024",
    "val_uah": "141 500",
    "val_words": "Сто сорок одна тисяча п'ятсот гривень 00 коп.",
    "head_pib": "Шаповал Т.Ю.",
    "head_full": "Шаповал Тимур Юрійович",
    "current_date": "14.05.2026",
    "amort_total": "51 883,26",
    "temp_reg_doc": "ХХР №087772"
}

output_dir = r"d:\ГО Талан UA\Talan UA Antigravity manager\_DROPZONE\OUT"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

def create_doc(name, title, content_lines):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for line in content_lines:
        p = doc.add_paragraph()
        if line.startswith("**") and line.endswith("**"):
            run = p.add_run(line.strip("**"))
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.add_run(line)
    
    file_path = os.path.join(output_dir, f"{name}.docx")
    doc.save(file_path)

# 1. Наказ
create_doc("1_Nakaz_Vvedennya", "НАКАЗ", [
    data["org_name"],
    f"(ЄДРПОУ {data['edrpou']})",
    "",
    "**НАКАЗ № 1-А**",
    f"**від «{data['reg_date'][:2]}» червня 2024 р.**",
    "**м. Черкаси**",
    "",
    "**Про зарахування на баланс та введення в експлуатацію гуманітарної допомоги (автомобіля)**",
    "",
    f"У зв’язку з отриманням тимчасового реєстраційного талону {data['temp_reg_doc']} від {data['reg_date']} та з метою забезпечення статутної діяльності ГО «Талан ЮА»,",
    "",
    "НАКАЗУЮ:",
    f"1. Зарахувати на баланс та ввести в експлуатацію з 13 червня 2024 року автомобіль {data['car_brand']}, рік випуску: {data['year']}, VIN: {data['vin']}, державний номер: {data['ua_plate']} (іноземний: {data['de_plate']}), отриманий як гуманітарна допомога згідно з Декларацією (Унікальний код: {data['unique_code']}) від {data['decl_date']}.",
    f"2. Визначити первісну вартість на рівні {data['val_uah']} грн ({data['val_words']}).",
    "3. Встановити строк корисного використання — 5 років (60 місяців).",
    f"4. Призначити відповідальною особою за експлуатацію {data['head_full']}.",
    "",
    f"Голова ГО «Талан ЮА» ________________ / {data['head_pib']} /"
])

# 2. Акт введення
create_doc("2_Akt_Vvedennya", "АКТ", [
    f"ЗАТВЕРДЖУЮ Голова ГО «Талан ЮА» ____________ / {data['head_pib']} /",
    f"«{data['reg_date'][:2]}» червня 2024 р.",
    "",
    "**АКТ № 1**",
    "**введення в експлуатацію транспортного засобу**",
    "",
    f"Комісія склала цей Акт про те, що автомобіль {data['car_brand']}, державний номер {data['ua_plate']} (VIN: {data['vin']}), введений в експлуатацію для статутної діяльності.",
    f"Підстава: Тимчасовий реєстраційний талон {data['temp_reg_doc']}.",
    "Технічний стан: справний.",
    "",
    "Голова комісії: ________________ / {data['head_pib']} /",
    "Член комісії: ________________ / [ПІБ] /"
])

# 3. Акт оцінки
create_doc("3_Akt_Otsinky", "АКТ ОЦІНКИ", [
    "**АКТ № 1-О**",
    "**оцінки справедливої вартості**",
    "",
    f"Комісія встановила справедливу вартість автомобіля {data['car_brand']}, номер {data['ua_plate']}, VIN: {data['vin']} станом на червень 2024 року.",
    f"Ринкова вартість визначена у розмірі: {data['val_uah']} грн ({data['val_words']}).",
    "",
    "Підписи комісії: ________________ / ________________"
])

# 4. Бух довідка
create_doc("4_Bukh_Dovidka", "ДОВІДКА", [
    "**БУХГАЛТЕРСЬКА ДОВІДКА № 1**",
    f"**від «{data['current_date'][:2]}» травня 2026 р.**",
    "",
    f"У зв'язку з технічною помилкою донараховано амортизацію за період 07.2024 - 04.2026 за автомобіль {data['car_brand']} (номер {data['ua_plate']}).",
    f"Сума донарахування: {data['amort_total']} грн.",
    "",
    f"Голова ГО ________________ / {data['head_pib']} /"
])

# 5. Лист
create_doc("5_Lyst_Poyasnennya", "ЛИСТ", [
    "Міністерству соціальної політики України",
    "",
    "**ЛИСТ-ПОЯСНЕННЯ**",
    "",
    f"ГО «Талан ЮА» (ЄДРПОУ {data['edrpou']}) повідомляє про цільове використання авто {data['car_brand']}, держ. номер {data['ua_plate']}, ввезеного за кодом {data['unique_code']}.",
    "Автомобіль вчасно поставлений на облік (Талон ХХР №087772 від 13.06.2024).",
    "Просимо розблокувати систему good.gov.ua.",
    "",
    f"Голова ГО ________________ / {data['head_pib']} /"
])
