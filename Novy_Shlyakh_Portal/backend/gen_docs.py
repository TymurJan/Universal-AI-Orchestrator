import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define the data
data = {
    "org_name": "ГРОМАДСЬКА ОРГАНІЗАЦІЯ «ТАЛАН ЮА»",
    "edrpou": "45119390",
    "address": "18005, Черкаська обл., м. Черкаси, вул. Волкова, буд. 95",
    "car_brand": "SKODA FABIA COMBI",
    "vin": "TMBGC26Y364517308",
    "number_plate": "GD MK826",
    "year": "2005",
    "unique_code": "23844937",
    "decl_date": "11.05.2024",
    "reg_date": "13.06.2024",
    "val_uah": "141 500",
    "val_words": "Сто сорок одна тисяча п'ятсот гривень 00 коп.",
    "head_pib": "Шаповал Т.Ю.",
    "head_full": "Шаповал Тимур Юрійович",
    "current_date": "12.05.2026",
    "amort_total": "51 883,26"
}

output_dir = r"C:\Users\style\.gemini\antigravity\brain\fbeb767b-b153-42b8-b0ad-f9a020659dee"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

def create_doc(name, title, content_lines):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for line in content_lines:
        p = doc.add_paragraph()
        if line.startswith("**") and line.endswith("**"):
            run = p.add_run(line.strip("**"))
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("ЗАТВЕРДЖУЮ") or line.startswith("Голова ГО"):
             p.add_run(line)
        else:
            p.add_run(line)
    
    file_path = os.path.join(output_dir, f"{name}.docx")
    doc.save(file_path)
    print(f"Saved: {file_path}")

# 1. Наказ
create_doc("1_Nakaz_Vvedennya", "НАКАЗ", [
    data["org_name"],
    f"(ЄДРПОУ {data['edrpou']})",
    "",
    "**НАКАЗ № 1-А**",
    f"**від «{data['reg_date'][:2]}» червня 2024 р.**",
    "**м. Черкаси**",
    "",
    "**Про зарахування на баланс та введення в експлуатацію гуманітарної допомоги (автомобіля)**",
    "",
    f"У зв’язку з отриманням свідоцтва про реєстрацію транспортного засобу (тимчасовий облік) від {data['reg_date']} та з метою забезпечення статутної діяльності ГО «Талан ЮА»,",
    "",
    "НАКАЗУЮ:",
    f"1. Зарахувати на баланс та ввести в експлуатацію з 13 червня 2024 року автомобіль {data['car_brand']}, рік випуску: {data['year']}, VIN: {data['vin']}, іноземний номер: {data['number_plate']}, отриманий згідно з Декларацією (Унікальний код: {data['unique_code']}) від {data['decl_date']}.",
    f"2. Визначити первісну вартість на рівні {data['val_uah']} грн ({data['val_words']}).",
    "3. Встановити строк корисного використання — 5 років (60 місяців).",
    f"4. Призначити відповідальною особою {data['head_full']}.",
    "",
    f"Голова ГО «Талан ЮА» ________________ / {data['head_pib']} /"
])

# 2. Акт введення
create_doc("2_Akt_Vvedennya", "АКТ", [
    f"ЗАТВЕРДЖУЮ Голова ГО «Талан ЮА» ____________ / {data['head_pib']} /",
    f"«{data['reg_date'][:2]}» червня 2024 р.",
    "",
    "**АКТ № 1**",
    "**введення в експлуатацію транспортного засобу**",
    "",
    f"Комісія склала цей Акт про те, що автомобіль {data['car_brand']}, VIN: {data['vin']}, введений в експлуатацію для статутної діяльності.",
    "Технічний стан: справний.",
    "Комплектація: ключі, техпаспорт.",
    "",
    "Голова комісії: ________________ / [ПІБ] /",
    "Член комісії: ________________ / [ПІБ] /"
])

# 3. Акт оцінки
create_doc("3_Akt_Otsinky", "АКТ ОЦІНКИ", [
    "**АКТ № 1-О**",
    "**оцінки справедливої вартості**",
    "",
    f"Комісія встановила справедливу вартість автомобіля {data['car_brand']}, VIN: {data['vin']} станом на червень 2024 року.",
    f"Ринкова вартість: {data['val_uah']} грн ({data['val_words']}).",
    "",
    "Підписи комісії: ________________ / ________________"
])

# 4. Бух довідка
create_doc("4_Bukh_Dovidka", "ДОВІДКА", [
    "**БУХГАЛТЕРСЬКА ДОВІДКА № 1**",
    f"**від «{data['current_date'][:2]}» травня 2026 р.**",
    "",
    f"Донараховано амортизацію за період 07.2024 - 04.2026 за автомобіль {data['car_brand']}.",
    f"Сума донарахування: {data['amort_total']} грн.",
    "",
    f"Голова ГО ________________ / {data['head_pib']} /"
])

# 5. Лист
create_doc("5_Lyst_Poyasnennya", "ЛИСТ", [
    "Міністерству соціальної політики України",
    "",
    "**ЛИСТ-ПОЯСНЕННЯ**",
    "",
    f"ГО «Талан ЮА» (ЄДРПОУ {data['edrpou']}) повідомляє про цільове використання авто {data['car_brand']}, ввезеного за кодом {data['unique_code']}.",
    "Просимо розблокувати систему good.gov.ua.",
    "",
    f"Голова ГО ________________ / {data['head_pib']} /"
])
