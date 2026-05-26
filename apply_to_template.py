from pathlib import Path
from docx import Document
import shutil
from datetime import datetime

TEMPLATE = Path(r"D:\ГО Талан UA\Гранти\Конкурс IREX\1. RFA_UVRR_Reginal_Coordination_Project Proposal_template.docx")
BACKUP = TEMPLATE.with_suffix('.docx.bak.' + datetime.now().strftime('%Y%m%d%H%M%S'))

# Texts to insert (from COMPLIANCE_MAP)
EXEC_SUM_FIRST = (
    'Створення інформаційно-координаційного ХАБу "Новий Шлях" допоможе ветерану не витрачати критично дефіцитні сили, '
    'час, емоції та здоров\'я на хаотичний пошук фахівців, записи на візити та оформлення документів, яких у нього (фізично та ментально) часто просто немає, '
    'особливо якщо є проблеми з пересуванням через ампутації тощо.'
)
ANALYTICAL_BRIEFS = (
    'Показник для Objective 4: Провести 5 аналітичних брифів — по одному звіту кожні 2 місяці після запуску для оперативного донесення виявлених прогалин до органів влади.'
)
BUDGET_NOTE = (
    'Власний нефінансовий внесок ГО «ТАЛАН ЮА» (In-kind contribution): $3,500. '
    'Запитуване фінансування від IREX: $20,000 (виділено в IT-статті бюджету для розгортання, безпеки та тестування).'
)

if not TEMPLATE.exists():
    print('Template not found at', TEMPLATE)
    raise SystemExit(1)

# Backup original
shutil.copy2(TEMPLATE, BACKUP)
print('Backed up template to', BACKUP)

# Load template
doc = Document(TEMPLATE)

# helper to find paragraph index containing any of a list of keys
def find_paragraph_index(keys):
    for i,p in enumerate(doc.paragraphs):
        for k in keys:
            if k in p.text:
                return i
    return None

# Insert executive summary first sentence after '1. Summary' heading
summary_keys = ['1. Summary', '1. Summary (up to 400 words)', 'Короткий опис', '1. Summary (up to 400 words)/ Короткий опис']
idx = find_paragraph_index(summary_keys)
if idx is not None:
    # find next non-empty paragraph after heading
    j = idx + 1
    while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() == '':
        j += 1
    if j < len(doc.paragraphs):
        para = doc.paragraphs[j]
        if EXEC_SUM_FIRST not in para.text:
            para.text = EXEC_SUM_FIRST + '\n\n' + para.text
            print('Inserted executive summary first sentence at paragraph', j+1)
    else:
        doc.add_paragraph(EXEC_SUM_FIRST)
        print('Appended executive summary at end')
else:
    print('Summary heading not found; appending executive summary at end')
    doc.add_paragraph(EXEC_SUM_FIRST)

# Ensure Objective 4 contains analytics brief requirement
obj_keys = ['Objective 4', '4.', 'Objective 4:']
obj_idx = find_paragraph_index(obj_keys)
if obj_idx is not None:
    if 'аналітич' not in doc.paragraphs[obj_idx].text:
        doc.paragraphs[obj_idx].text = doc.paragraphs[obj_idx].text + ' ' + ANALYTICAL_BRIEFS
        print('Updated Objective 4 at paragraph', obj_idx+1)
else:
    print('Objective 4 heading not found; appending analytics brief at end')
    doc.add_paragraph(ANALYTICAL_BRIEFS)

# Insert budget note near 'Budget' or '15.' headings
budget_keys = ['15.', 'Budget', 'Бюджет']
bidx = find_paragraph_index(budget_keys)
if bidx is not None:
    # insert after the heading's next non-empty paragraph
    j = bidx + 1
    while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() == '':
        j += 1
    # insert a new paragraph before j (if j==len -> append)
    if j < len(doc.paragraphs):
        # insert by creating new paragraph at position j
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        # fallback: append after j
        doc.paragraphs[j-1]._parent.add_paragraph(BUDGET_NOTE)
        print('Inserted budget note after paragraph', j)
    else:
        doc.add_paragraph(BUDGET_NOTE)
        print('Appended budget note at end')
else:
    doc.add_paragraph(BUDGET_NOTE)
    print('Budget heading not found; appended budget note at end')

# Save changes back to original TEMPLATE (overwrite)
doc.save(TEMPLATE)
print('Saved filled template to', TEMPLATE)
print('Done')
