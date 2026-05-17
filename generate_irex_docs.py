import os
import sys
from pathlib import Path
import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = Path(r"D:\ГО Талан UA\Talan UA Antigravity manager\_DROPZONE\OUT")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# --- ДСТУ 4163:2020 & Talan Legal Skill Standards ---
def create_dstu_document():
    doc = docx.Document()
    # Поля: Ліве — 30 мм, Праве — 10 мм, Верхнє/Нижнє — 20 мм
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
    
    # Базовий шрифт: Times New Roman, 11 pt (еталон для ГО)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    return doc

def add_talan_letterhead(doc):
    """Фірмовий Бланк (Шапка) ГО 'Талан ЮА' за суворими правилами скіла"""
    # Назва Організації: Шрифт 14, не жирний, space_after = 8pt
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("ГРОМАДСЬКА ОРГАНІЗАЦІЯ «ТАЛАН ЮА»")
    run_org.font.size = Pt(14)
    run_org.font.bold = False
    p_org.paragraph_format.space_after = Pt(8)
    p_org.paragraph_format.line_spacing = 1.0
    
    # Реквізити банку та ЄДРПОУ: Шрифт 8, space_after = 8pt
    p_req = doc.add_paragraph()
    p_req.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_req = p_req.add_run("ЄДРПОУ: 45119390 | р/р UA49305299000002600103160858113 у АТ КБ «ПРИВАТБАНК», МФО 305299")
    run_req.font.size = Pt(8)
    p_req.paragraph_format.space_after = Pt(8)
    p_req.paragraph_format.line_spacing = 1.0
    
    # Адреса та контакти: Шрифт 8, space_after = 16pt
    p_addr = doc.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_addr = p_addr.add_run("Юридична адреса: 18005, м. Черкаси, вул. Волкова, буд. 95, кв. 26 | Тел: +38 (063) 656-89-99 | ngo.talan.ua@gmail.com")
    run_addr.font.size = Pt(8)
    p_addr.paragraph_format.space_after = Pt(16)
    p_addr.paragraph_format.line_spacing = 1.0

def add_header_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)
    
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(12)
        run2.font.italic = True
        p2.paragraph_format.space_after = Pt(16)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

def add_bullet(doc, bold_prefix, text, is_last_before_signature=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run_bold = p.add_run(bold_prefix)
    run_bold.bold = True
    p.add_run(text)
    if is_last_before_signature:
        # Алгоритм захисту від порожньої сторінки: останній абзац тримається з підписами
        p.paragraph_format.keep_with_next = True

def add_paragraph_with_keep(doc, text, is_last_before_signature=False):
    p = doc.add_paragraph(text)
    if is_last_before_signature:
        # Алгоритм захисту від порожньої сторінки: останній абзац тримається з підписами
        p.paragraph_format.keep_with_next = True
    return p

def add_table_row(table, data, is_header=False):
    row = table.add_row()
    for i, text in enumerate(data):
        cell = row.cells[i]
        cell.text = str(text)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if is_header:
            for run in p.runs:
                run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if i == 0 or (len(data) > 2 and i == len(data)-1):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return row

def add_talan_bilateral_signature(doc, partner_org, partner_title, partner_name):
    """
    Сувора верстка підписів за ДСТУ та правилами Talan Legal Skill:
    - Заголовок «Підписи сторін:» звичайним (не жирним) шрифтом, центрується, keep_with_next=True.
    - Сторона 1 (ГО «ТАЛАН ЮА») — ЛІВИЙ край.
    - Сторона 2 (Партнер) — ПРАВИЙ край.
    - Прізвища ВЕЛИКИМИ ЛІТЕРАМИ.
    """
    p_sig_head = doc.add_paragraph()
    p_sig_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_head.paragraph_format.space_before = Pt(24)
    p_sig_head.paragraph_format.space_after = Pt(12)
    p_sig_head.paragraph_format.keep_with_next = True
    run_sig_head = p_sig_head.add_run("Підписи сторін:")
    run_sig_head.font.bold = False # Суворо звичайним шрифтом
    
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(90) # Ліва сторона
    table.columns[1].width = Mm(80) # Права сторона
    
    # Рядок 1: Посади та Назви Організацій
    cell_l0 = table.cell(0, 0)
    cell_l0.text = "Сторона 1\nГолова правління\nГО «ТАЛАН ЮА»"
    cell_l0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_l0.paragraphs[0].runs[0].font.bold = True
    
    cell_r0 = table.cell(0, 1)
    cell_r0.text = f"Сторона 2\n{partner_title}\n{partner_org}"
    cell_r0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT # Суворо по правому краю
    cell_r0.paragraphs[0].runs[0].font.bold = True
    
    # Рядок 2: Місце для підпису та ПІБ (Прізвище ВЕЛИКИМИ ЛІТЕРАМИ)
    cell_l1 = table.cell(1, 0)
    cell_l1.text = "\n_________________ / Тимур ШАПОВАЛ /\nМ.П."
    cell_l1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    cell_r1 = table.cell(1, 1)
    cell_r1.text = f"\n_________________ / {partner_name} /\nМ.П."
    cell_r1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_single_signature_block(doc, org_name, signer_title, signer_name):
    """Для односторонніх листів підтримки від партнерів до IREX"""
    p_sig_head = doc.add_paragraph()
    p_sig_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_head.paragraph_format.space_before = Pt(24)
    p_sig_head.paragraph_format.space_after = Pt(12)
    p_sig_head.paragraph_format.keep_with_next = True
    run_sig_head = p_sig_head.add_run("Підпис сторони:")
    run_sig_head.font.bold = False # Суворо звичайним шрифтом
    
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(170)
    
    cell_0 = table.cell(0, 0)
    cell_0.text = f"{signer_title}\n{org_name}"
    cell_0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_0.paragraphs[0].runs[0].font.bold = True
    
    cell_1 = table.cell(1, 0)
    cell_1.text = f"\n_________________ / {signer_name} /\nМ.П."
    cell_1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ==========================================
# 1. GENERATE MAIN GRANT PROPOSAL
# ==========================================
def generate_main_proposal():
    doc = create_dstu_document()
    add_talan_letterhead(doc)
    add_header_title(doc, "ГРАНТОВА ЗАЯВКА НА КОНКУРС IREX", "Проєкт: «Новий Шлях: Цифрова диспетчерська та регіональний координаційний хаб підтримки ветеранів Черкащини»")
    
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ["Параметр", "Дані заявки"], is_header=True)
    add_table_row(table, ["Конкурс IREX", "Регіональна координація систем підтримки ветеранів (UVRR RFA)"])
    add_table_row(table, ["Організація-заявник", "Громадська організація «ТАЛАН ЮА» (ЄДРПОУ: 45119390)"])
    add_table_row(table, ["Сума запиту / Тривалість", "$100,000 USD / 12 місяців"])
    add_table_row(table, ["Географія проєкту", "Черкаська область (м. Черкаси та територіальні громади)"])
    add_table_row(table, ["Веб-портал / Бот", "https://novyshlyakh.ua | @Veteran_NovyShlyakh_Bot"])
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_section_heading(doc, "1. Анотація проєкту (Executive Summary)")
    doc.add_paragraph("Проєкт «Новий Шлях» створює інноваційну гібридну модель регіональної координації ветеранських послуг у Черкаській області. Замість традиційних бюрократичних рад та статичних довідників, ми впроваджуємо «Цифрову диспетчерську» — platform прямої дії, що об’єднує верифіковану спільноту фахівців (юристів, психологів, реабілітологів, роботодавців) та ветеранську спільноту.")
    doc.add_paragraph("Проєкт забезпечує найкоротший шлях від потреби ветерана до реальної послуги через автоматизовану систему AI-маршрутизації, обов'язковий контроль якості (follow-up за 48 годин) та фізичну точку супроводу (Intake Point) у Черкасах. Наша мета — перетворити хаотичний ринок допомоги на керовану, прозору та підзвітну екосистему.")

    add_section_heading(doc, "2. Обґрунтування актуальності (Problem Statement)")
    doc.add_paragraph("Ветеранська спільнота Черкащини стикається з трьома критичними бар’єрами при поверненні до цивільного життя:")
    add_bullet(doc, "1. Інформаційний хаос та фрагментарність: ", "Послуги розпорошені між державними установами, різними ГО та приватним сектором. Ветеран витрачає місяці на пошук потрібного фахівця.")
    add_bullet(doc, "2. Низька довіра та відсутність фільтрів якості: ", "Відсутня інституція, яка б перевіряла реальну кваліфікацію юристів чи психологів, що пропонують послуги ветеранам.")
    add_bullet(doc, "3. Бюрократична затримка: ", "Довгий шлях від первинного звернення до отримання конкретної допомоги (оформлення УБД, проходження МСЕК/ВЛК, компенсації).")
    doc.add_paragraph("Існуючі координаційні ради на рівні області часто фокусуються на стратегічних та протокальних нарадах, залишаючи щоденну операційну взаємодію «ветеран — спеціаліст» некерованою. Проєкт «Новий Шлях» вирішує саме цю проблему, надаючи владі та громаді дієвий цифровий інструмент прямої координації.")

    add_section_heading(doc, "3. Мета та завдання проєкту (Goal & Objectives)")
    doc.add_paragraph("Мета: Побудувати сталу регіональну екосистему підтримки та реінтеграції ветеранів через об'єднання ресурсів громадського, державного та приватного секторів Черкащини в єдину «Цифрову диспетчерську».")
    add_bullet(doc, "Objective 1: ", "Сформувати та верифікувати регіональну мережу з 150+ фахівців Черкащини (юристи, кризові психологи, кар'єрні радники, лікарі) з інтегрованою системою рейтингування.")
    add_bullet(doc, "Objective 2: ", "Модернізувати IT-інфраструктуру порталу «Новий Шлях» (Telegram Mini App + Web) до рівня захищеного автоматизованого диспетчерського центру на базі професійної СУБД SQLite з логуванням усіх звернень.")
    add_bullet(doc, "Objective 3: ", "Забезпечити регулярний фізичний супровід та консультування ветеранів через офлайн-хаб (Intake Point) у м. Черкаси для проведення 1-2 глибинних прийомів щодня.")
    add_bullet(doc, "Objective 4: ", "Провести 5 стратегічних координаційних заходів (круглі столи, панельні дискусії, прес-конференції) для повної інтеграції платформи в систему Черкаської ОВА, міської ради та територіальних громад.")

    add_section_heading(doc, "4. Досвід організації (Past Performance & Organizational Capacity)")
    doc.add_paragraph("ГО «Талан ЮА» має підтверджену історію безперервної кризової, виробничої та логістичної роботи, починаючи з другого дня повномасштабного вторгнення РФ. Хоча для організації це перша масштабна інституційна грантова заявка, команда володіє потужним практичним досвідом управління складними проєктними циклами в умовах високого ризику та обмежених ресурсів.")
    add_bullet(doc, "Виробництво та постачання тактичного спорядження: ", "Організація власними силами з нуля налагодила виробництво та безперебійне постачання розвантажувальних систем, РПС, спальних мішків, матраців, тактичних рукавиць, устілок, бафів та балаклав для передових підрозділів Сил оборони.")
    add_bullet(doc, "Складна технічна логістика та ремонт: ", "Команда забезпечувала не лише закупівлю та доставку десятків автомобілів, шин, FPV-дронів, систем Starlink, маскувальних сіток та спеціалізованого оборонного устаткування, а й організувала системний ремонт і відновлення як легкового транспорту, так і важкої військової техніки.")
    add_bullet(doc, "Операційний менеджмент (Кожен збір як проєктний цикл): ", "Кожну волонтерську кампанію команда реалізовувала за класичною методологією проєктного менеджменту: ідентифікація проблеми → збір фінансових ресурсів → пошук та експертна закупівля → складна логістика → повна фінансова та публічна звітність.")

    add_section_heading(doc, "5. Запропонований координаційний механізм (Coordination Mechanism)")
    doc.add_paragraph("Ми пропонуємо модель «Децентралізованої цифрової мережі координації прямої дії».")
    add_bullet(doc, "Мандат: ", "Координація надавачів послуг у громадах та жорстка верифікація їхньої професійної етики й якості.")
    add_bullet(doc, "Склад екосистеми: ", "ГО «Талан ЮА» (адміністратор та IT-архітектор), Черкаська ОВА та міська рада (стратегічні стейкхолдери), Мережа спеціалістів (безпосередні надавачі послуг).")
    add_bullet(doc, "Процес взаємодії: ", "Ветеран подає запит через Telegram-бот або веб-портал. AI-система миттєво маршрутизує його до найближчого вільного фахівця потрібного профілю. Диспетчер ГО фіксує звернення в базі та контролює виконання запиту (автоматизований follow-up фідбек через 48 годин).")

    add_section_heading(doc, "6. Команда та управління проєктом (Key Personnel & Management Structure)")
    doc.add_paragraph("Управління проєктом здійснюватиметься збалансованою командою лідерів із підтвердженим досвідом у громадському секторі та сфері комунікацій:")
    add_bullet(doc, "Керівник комунікаційного напрямку (Communications & PR Lead): Сакун Анна Валеріївна. ", "Фахівчиня з багаторічним успішним досвідом у сфері публічних комунікацій та зв'язків із громадськістю. Анна разом зі своєю командою є лауреаткою престижної премії Фонду родини Богдана Гаврилишина (Bohdan Hawrylyshyn Family Foundation). Цей статус підтверджує її високий професіоналізм, лідерські якості, бездоганну репутацію та здатність будувати ефективні, довірливі стосунки з ветеранською спільнотою та медіа.")
    add_bullet(doc, "Керівник проєкту (Project Lead) та Фінансовий менеджер: ", "Наразі ГО «Талан ЮА» сформувала шорт-лист висококваліфікованих кандидатів на ключові адміністративні та фінансові посади. Фінальний відбір проводиться з урахуванням жорстких кваліфікаційних вимог IREX щодо досвіду управління міжнародними грантами, комплаєнсу та звітності.")
    add_bullet(doc, "Технічний архітектор (IT & AI Lead): ", "Забезпечує безперебійну роботу portal, захист бази даних SQLite, налаштування серверної інфраструктури та алгоритмів маршрутизації.")

    add_section_heading(doc, "7. Стратегічні партнерства (Strategic Partnerships & Stakeholder Engagement)")
    doc.add_paragraph("Для забезпечення максимального охоплення та ефективної маршрутизації ветеранів, ГО «Талан ЮА» вибудовує потужну міжсекторальну коаліцію на рівні Черкаської області. Наразі досягнуто стратегічних домовленостей про partnership та інформаційний обмін із ключовими стейкхолдерами:")
    add_bullet(doc, "Органи місцевого самоврядування та влади: ", "Департамент соціальної політики Черкаської міської ради (забезпечення безперешкодного доступу до державних соціальних послуг та пряма взаємодія з мережею ЦНАПів).")
    add_bullet(doc, "Сектор безпеки та оборони: ", "Добровольче формування Черкаської територіальної громади (ДФ Черкаської ТГ) — взаємодія з чинними бійцями, ветеранами та організація підтримки на етапі демобілізації (Pre-Exit Support).")
    add_bullet(doc, "Провідні громадські організації та ветеранські простори: ", "ГО «Черкаський інститут міста», ГО «Горизонт змін», а також спеціалізований ветеранський простір «Ветеран Про». Це partnership гарантує гармонійну інтеграцію platform «Новий Шлях» у діючу громадську екосистему Черкащини без дублювання функцій та створює потужну синергію.")

    add_section_heading(doc, "8. Стійкість та стратегія виходу (Sustainability)")
    add_bullet(doc, "Технічна сталість: ", "Платформа розроблена на легких, оптимізованих технологіях (SQLite, Python, Telegram API, Vanilla Web), що зводить щомісячні витрати на сервери та підтримку до абсолютного мінімуму.")
    add_bullet(doc, "Інституціоналізація: ", "Підписання меморандумів (MoUs) з ОТГ Черкащини про офіційне використання «Нового Шляху» як головного реєстру надавачів ветеранських послуг у громадах.")
    add_bullet(doc, "Соціальна та фінансова сталість: ", "Спільнота верифікованих спеціалістів працює на умовах квоти «25% соціального внеску» (безкоштовні або пільгові послуги для ветеранів в обмін на доступ до клієнтської web-бази portal). Це створює постійний, автономний обіг експертних ресурсів всередині екосистеми.")

    add_section_heading(doc, "9. Комунікації та розповсюдження (Communication Strategy)")
    add_bullet(doc, "Рішення 1 (Офлайн-навігація): ", "Кампанія «Прямий доступ» — розміщення 100+ високоякісних плакатів формату А3 із прямим QR-кодом на бот у всіх ЦНАПах, лікарнях, ТЦК та державних установах Черкаської області. Розповсюдження 1000+ інформаційних пам'яток та встановлення мобільних стендів (павуків).")
    add_bullet(doc, "Рішення 2 (Digital-охоплення): ", "Запуск потужної таргетованої реклами у Facebook та Instagram, спрямованої безпосередньо на ветеранів, військовослужбовців та членів їхніх родин на Черкащині.")
    add_bullet(doc, "Рішення 3 (Публічний PR): ", "Організація круглих столів та прес-конференцій із залученням регіональних ЗМІ для презентації досягнень platform та легітимізації спільноти.")

    add_section_heading(doc, "10. Орієнтовний бюджет ($100,000) — Розподіл за статтями")
    btable = doc.add_table(rows=1, cols=3, style='Table Grid')
    btable.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(btable, ["Стаття витрат", "Бюджет (USD)", "Частка (%)"], is_header=True)
    add_table_row(btable, ["Персонал (Менеджмент, Комунікаційний лід, Диспетчери хабу)", "$35,000", "35%"])
    add_table_row(btable, ["IT-розробка, СУБД SQLite, Захист даних та AI-модуль", "$20,000", "20%"])
    add_table_row(btable, ["Оренда, облаштування та утримання фізичного Intake-хабу в Черкасах", "$15,000", "15%"])
    add_table_row(btable, ["PR, Маркетинг, Друк плакатів А3 (100 шт.), Листівки (1000 шт.), Павуки, SMM", "$15,000", "15%"])
    add_table_row(btable, ["Публічні заходи (Круглі столи, Оренда залів, Прес-конференції)", "$10,000", "10%"])
    add_table_row(btable, ["Адміністративні витрати (Бухгалтерський супровід, Комплаєнс, Зв'язок)", "$5,000", "5%"])
    add_table_row(btable, ["РАЗОМ:", "$100,000", "100%"])
    
    # Алгоритм захисту від порожньої сторінки для підписів
    p_last = doc.add_paragraph("Заявка складена та підтверджена командою ГО «Талан ЮА».")
    p_last.paragraph_format.keep_with_next = True
    
    add_talan_bilateral_signature(doc, "IREX UVRR", "Конкурсна комісія", "ПРЕДСТАВНИК IREX")
    
    doc.save(OUT_DIR / "01_IREX_Grant_Proposal_NovyShlyakh_FULL_v4.docx")
    print("✅ Main proposal generated (v4).")

# ==========================================
# 2. GENERATE LETTERS OF SUPPORT
# ==========================================
def generate_support_letter(filename, partner_name, partner_title, partner_org, specific_text):
    doc = create_dstu_document()
    
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_head = p_head.add_run("Міжнародній раді наукових досліджень та обмінів (IREX)\nКонкурсна комісія програми підтримки ветеранів\n\n")
    run_head.font.bold = True
    
    p_from = doc.add_paragraph()
    p_from.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_from.add_run(f"від: {partner_org}\nАдреса: м. Черкаси\n").font.italic = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ЛИСТ ПІДТРИМКИ ПРОЄКТУ «НОВИЙ ШЛЯХ»")
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    p_title.paragraph_format.space_after = Pt(18)
    
    doc.add_paragraph(f"{partner_org} засвідчує свою повагу Міжнародній раді наукових досліджень та обмінів (IREX) та цим листом висловлює цілковиту підтримку проєктній заявці Громадської організації «ТАЛАН ЮА» (ЄДРПОУ: 45119390) під назвою «Новий Шлях: Цифрова диспетчерська та регіональний координаційний хаб підтримки ветеранів Черкащини».")
    
    doc.add_paragraph(specific_text)
    
    doc.add_paragraph("Ми переконані, що створення гібридної «Цифрової диспетчерської» є нагальною та критично важливою ініціативою для Черкаської області. Платформа «Новий Шлях» дозволить об'єднати зусилля державного, громадського та приватного секторів, усунути бюрократичні бар'єри та забезпечити ветеранам швидкий і прозорий доступ до верифікованих фахівців.")
    
    doc.add_paragraph(f"Зі свого боку, {partner_org} підтверджує готовність до активної співпраці з ГО «Талан ЮА», участі в роботі регіональної координаційної ради, спільного проведення інформаційних заходів та підписання відповідного Меморандуму про співпрацю (MoU) після початку реалізації проєкту.")
    
    # Алгоритм захисту від порожньої сторінки: keep_with_next на останньому абзаці
    p_last = add_paragraph_with_keep(doc, "Просимо конкурсну комісію IREX підтримати цю ініціативу, яка має стратегічне значення для успішної реінтеграції ветеранів та членів їхніх родин у нашій громаді.", is_last_before_signature=True)
    
    # Верстка підпису за правилами скіла
    add_single_signature_block(doc, partner_org, partner_title, partner_name)
    
    doc.save(OUT_DIR / filename)
    print(f"✅ Support letter generated: {filename}")

partners = [
    (
        "02_Letter_of_Support_Cherkasy_City_Council_v4.docx",
        "ІМ'Я ПРІЗВИЩЕ", 
        "Директор Департаменту",
        "Департамент соціальної політики Черкаської міської ради",
        "Враховуючи актуальність питань соціального захисту та адаптації демобілізованих військовослужбовців, Департамент соціальної політики відзначає високий потенціал цифрової платформи «Новий Шлях». Впровадження автоматизованої системи маршрутизації та облік звернень через Intake-хаб ідеально доповнюють існуючу державну інфраструктуру та забезпечать ефективну взаємодію з мережею ЦНАПів та управліннями соціального захисту населення Черкаської міської ради."
    ),
    (
        "03_Letter_of_Support_DF_Cherkasy_TG_v4.docx",
        "ІМ'Я ПРІЗВИЩЕ",
        "Командир",
        "Добровольче формування Черкаської територіальної громади (ДФ Черкаської ТГ)",
        "Командування ДФ Черкаської ТГ особливо відзначає практичну спрямованість ініціативи ГО «Талан ЮА». Для наших бійців, які проходять етап демобілізації або потребують фахової підтримки (Pre-Exit Support), наявність єдиного цифрового вікна з перевіреними юристами, психологами та лікарями є життєво необхідною. Ми готові виступати містком між platform та чинними військовослужбовцями для забезпечення їхньої безшовної адаптації."
    ),
    (
        "04_Letter_of_Support_NGO_Cherkasy_City_Institute_v4.docx",
        "ІМ'Я ПРІЗВИЩЕ",
        "Голова правління",
        "Громадська організація «Черкаський інститут міста»",
        "Як організація, що займається міським розвитком та аналітикою громадських просторів, ми бачимо величезну синергію у співпраці з ГО «Талан ЮА». Створення фізичного Intake-хабу та розгортання цифрового маркетплейсу послуг для ветеранів сприятиме формуванню інклюзивного та безпечного міського середовища в Черкасах. Ми готові надавати експертну підтримку в проведені спільних круглих столів та інтеграції проєкту в стратегію розвитку міста."
    ),
    (
        "05_Letter_of_Support_NGO_Horizon_of_Change_v4.docx",
        "ІМ'Я ПРІЗВИЩЕ",
        "Голова організації",
        "Громадська організація «Горизонт змін»",
        "ГО «Горизонт змін» активно підтримує ініціативи, спрямовані на посилення спроможності громад та розвиток людського капіталу. Проєкт «Новий Шлях» є яскравим прикладом інноваційного підходу до вирішення складних соціальних викликів. Об'єднання понад 150 фахівців області в єдину верифіковану мережу створить потужний прецедент міжсекторальної взаємодії, і ми готові долучатися до навчальних та адвокаційних заходів проєкту."
    ),
    (
        "06_Letter_of_Support_Veteran_Pro_v4.docx",
        "ІМ'Я ПРІЗВИЩЕ",
        "Керівник простору",
        "Ветеранський простір «Ветеран Про»",
        "Як діючий ветеранський простір, ми щодня стикаємося з потребою у швидкій та якісній переадресації складних запитів ветеранів. Платформа «Новий Шлях», ініційована ГО «Талан ЮА», надає саме той цифровий диспетчерський інструмент, якого бракує офлайн-хабам. Ми готові об'єднувати наші зусилля за принципом «рівний — рівному», спільно вести кейс-менеджмент та надавати наш майданчик для проведення консультацій."
    )
]

# ==========================================
# 3. GENERATE BUDGET ESTIMATE DOC
# ==========================================
def generate_budget_doc():
    doc = create_dstu_document()
    add_talan_letterhead(doc)
    add_header_title(doc, "ДЕТАЛЬНИЙ КОШТОРИС ПРОЄКТУ «НОВИЙ ШЛЯХ»", "Загальний бюджет: $100,000 USD (Конкурс IREX UVRR)")
    
    doc.add_paragraph("Цей документ містить розширену розшифровку статей витрат для внутрішнього аналізу командою ГО «Талан ЮА» перед перенесенням в офіційну Excel-форму IREX.").paragraph_format.space_after = Pt(16)
    
    table = doc.add_table(rows=1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ["Категорія витрат / Стаття", "Опис та призначення", "Сума (USD)", "%"], is_header=True)
    
    add_table_row(table, [
        "1. ПЕРСОНАЛ ПРОЄКТУ",
        "Оплата праці ключової команди протягом 12 місяців: Керівник проєкту ($1,500/міс), Комунікаційний лід Анна Сакун ($1,000/міс), 2 Диспетчери-координатори Intake-хабу ($1,000/міс сумарно).",
        "$35,000", "35%"
    ])
    add_table_row(table, [
        "2. IT-РОЗРОБКА ТА AI",
        "Модернізація Telegram Mini App та веб-порталу, перехід на захищену СУБД SQLite, розробка AI-модуля маршрутизації, налаштування серверів та системи автоматичного логування Intake/Feedback.",
        "$20,000", "20%"
    ])
    add_table_row(table, [
        "3. ОФЛАЙН INTAKE-ХАБ",
        "Оренда приміщення в центрі м. Черкаси (12 міс), косметичний ремонт під вимоги інклюзивності, закупівля меблів, комп'ютерної техніки, зарядної станції (EcoFlow) та поточне утримання.",
        "$15,000", "15%"
    ])
    add_table_row(table, [
        "4. PR ТА ПОЛІГРАФІЯ",
        "Друк 100+ плакатів А3 із QR-кодами для ЦНАПів та ТЦК, 1000+ буклетів/пам'яток, виготовлення 5 мобільних стендів (павуків), таргетована SMM-реклама FB/Insta на ветеранів області.",
        "$15,000", "15%"
    ])
    add_table_row(table, [
        "5. ПУБЛІЧНІ ЗАХОДИ",
        "Організація 5 стратегічних круглих столів, панельних дискусій та прес-конференцій (оренда залів, кава-брейки, роздаткові матеріали, трансфер учасників із громад області).",
        "$10,000", "10%"
    ])
    add_table_row(table, [
        "6. АДМІНІСТРАТИВНІ",
        "Бухгалтерський супровід проєкту (комплаєнс за стандартами IREX/USAID), банківські комісії, поштові витрати, канцелярія, інтернет та корпоративний мобільний зв'язок.",
        "$5,000", "5%"
    ])
    add_table_row(table, ["РАЗОМ ЗАЯВКА:", "Повний обсяг фінансування від IREX", "$100,000", "100%"])
    
    p = table.rows[-1].cells[0].paragraphs[0]
    p.runs[0].font.bold = True
    p2 = table.rows[-1].cells[2].paragraphs[0]
    p2.runs[0].font.bold = True
    
    # Алгоритм захисту від порожньої сторінки
    p_last = doc.add_paragraph("Кошторис перевірено та затверджено фінансовим відділом ГО «Талан ЮА».")
    p_last.paragraph_format.keep_with_next = True
    
    add_talan_bilateral_signature(doc, "IREX UVRR", "Фінансовий офіцер", "ПРЕДСТАВНИК IREX")
    
    doc.save(OUT_DIR / "07_IREX_Budget_Estimate_NovyShlyakh_v4.docx")
    print("✅ Budget doc generated (v4).")

if __name__ == "__main__":
    print("🚀 Starting generation of IREX Grant Package in _DROPZONE/OUT (v4 - ДСТУ 4163:2020)...")
    generate_main_proposal()
    for p in partners:
        generate_support_letter(p[0], p[1], p[2], p[3], p[4])
    generate_budget_doc()
    print("🎉 ALL DOCUMENTS GENERATED SUCCESSFULLY (v4)!")
