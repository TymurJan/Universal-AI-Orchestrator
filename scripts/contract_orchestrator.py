import os
import json
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Налаштування шляхів ---
BASE_DIR = Path(__file__).resolve().parent.parent.parent
LOG_FILE = BASE_DIR / "logs" / "b2b_ledger.json"
STORAGE_DIR = BASE_DIR / "Universal-AI-Orchestrator" / "data" / "b2b" / "contracts"

# Створюємо директорії, якщо їх немає
STORAGE_DIR.mkdir(parents=True, exist_ok=True)
(BASE_DIR / "logs").mkdir(parents=True, exist_ok=True)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class ContractOrchestrator:
    def __init__(self):
        self.ledger_path = LOG_FILE
        self._init_ledger()

    def _init_ledger(self):
        if not self.ledger_path.exists():
            with open(self.ledger_path, 'w', encoding='utf-8') as f:
                json.dump([], f)

    def verify_edrpou(self, edrpou):
        """
        В реальності тут буде запит до OpenDataBot або WebScraping.
        Поки що імітуємо пошук.
        """
        logging.info(f"ШІ верифікує ЄДРПОУ: {edrpou}")
        # Імітація знайдених даних
        mock_data = {
            "12345678": {"name": "ТОВ 'ПЕРЕМОГА ІНВЕСТ'", "boss": "Олександр ВЕЛИКИЙ"},
            "87654321": {"name": "ГРОМАДСЬКА СПІЛКА 'КІБЕР-ЩИТ'", "boss": "Марія ПРИЙМАЧЕНКО"}
        }
        return mock_data.get(edrpou, {"name": f"Юридична особа ({edrpou})", "boss": "Керівник компанії"})

    def log_request(self, client_data):
        """Фіксація запиту для звітності ГО."""
        try:
            with open(self.ledger_path, 'r', encoding='utf-8') as f:
                ledger = json.load(f)
            
            entry = {
                "timestamp": datetime.now().isoformat(),
                "client_email": client_data.get("email"),
                "edrpou": client_data.get("edrpou"),
                "company_name": client_data.get("company_name"),
                "tier": client_data.get("tier"),
                "status": "Generated & Sent"
            }
            ledger.append(entry)
            
            with open(self.ledger_path, 'w', encoding='utf-8') as f:
                json.dump(ledger, f, ensure_ascii=False, indent=4)
            logging.info(f"Запит зафіксовано у леджері: {client_data['email']}")
        except Exception as e:
            logging.error(f"Помилка логування: {e}")

    def generate_contract(self, client_data):
        """Генерація договору за ДСТУ 4163:2020."""
        doc = Document()
        
        # Налаштування полів (ДСТУ 4163:2020)
        sections = doc.sections
        for section in sections:
            section.left_margin = Mm(30)
            section.right_margin = Mm(10)
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)

        # Шапка ГО "Талан ЮА"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("ГРОМАДСЬКА ОРГАНІЗАЦІЯ «ТАЛАН ЮА»")
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        doc.add_paragraph("ЄДРПОУ 44903160\nwww.talanai.com\nngo.talan.ua@gmail.com").alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"\nДОГОВІР № {datetime.now().strftime('%Y%m%d')}-SaaS")
        run.bold = True
        run.font.size = Pt(12)

        # Текст (спрощений для демонстрації)
        p = doc.add_paragraph()
        company = client_data.get('company_name', 'КЛІЄНТ')
        p.add_run(f"Цей договір укладено між ГО «ТАЛАН ЮА» (Виконавець) та {company} (Замовник) про надання доступу до ПЗ 'Universal AI Orchestrator' (Тариф: {client_data['tier']}).")
        
        # Підписи
        doc.add_paragraph("\n\nПІДПИСИ СТОРІН:").alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        cells = table.rows[0].cells
        cells[0].text = "ВІД ВИКОНАВЦЯ:\n\n________________\nТимур ШАПОВАЛ"
        cells[1].text = f"ВІД ЗАМОВНИКА:\n\n________________\n{client_data.get('boss', '________________')}"

        filename = f"Contract_{client_data['edrpou']}_{datetime.now().strftime('%H%M%S')}.docx"
        filepath = STORAGE_DIR / filename
        doc.save(filepath)
        return str(filepath)

    def process_request(self, email, edrpou, tier, position):
        # 1. Верифікація
        info = self.verify_edrpou(edrpou)
        client_data = {
            "email": email,
            "edrpou": edrpou,
            "tier": tier,
            "company_name": info['name'],
            "boss": info['boss'],
            "position": position
        }

        # 2. Логування
        self.log_request(client_data)

        # 3. Генерація
        path = self.generate_contract(client_data)
        
        logging.info(f"Договір створено: {path}")
        return path

if __name__ == "__main__":
    # Тестовий запуск
    orch = ContractOrchestrator()
    orch.process_request("test@client.com", "12345678", "Corporate", "Директор")
