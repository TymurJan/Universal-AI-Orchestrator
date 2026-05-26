from pathlib import Path
from docx import Document

# Paths
INPUT = Path(r"D:\ГО Талан UA\Гранти\Конкурс IREX\Робоча\Заявка\01_IREX_Grant_Proposal_NovyShlyakh_FULL_v7.docx")
OUTPUT = Path(r"D:\ГО Талан UA\Гранти\Конкурс IREX\Робоча\Заявка\01_IREX_Grant_Proposal_NovyShlyakh_FULL_v7_patched.docx")

# Texts from COMPLIANCE_MAP_IREX_20260525.md
EXEC_SUM_FIRST = (
    '«Створення інформаційно-координаційного ХАБу "Новий Шлях" допоможе ветерану '
    'не витрачати критично дефіцитні сили, час, емоції та здоров\'я на хаотичний пошук фахівців, '
    'записи на візити та оформлення документів, яких у нього (фізично та ментально) часто просто немає, '
    'особливо якщо є проблеми з пересуванням через ампутації тощо.»'
)

ANALYTICAL_BRIEFS = (
    'Показник для Objective 4: Провести 5 аналітичних брифів — по одному звіту кожні 2 місяці '
    'після запуску для оперативного донесення виявлених прогалин до органів влади.'
)

BUDGET_NOTE = (
    'Власний нефінансовий внесок ГО «ТАЛАН ЮА» (In-kind contribution): $3,500. '
    'Запитуване фінансування від IREX: $20,000 (виділено в IT-статті бюджету для розгортання, безпеки та тестування).'
)

# Load document
if not INPUT.exists():
    print('INPUT not found:', INPUT)
    raise SystemExit(1)

doc = Document(INPUT)

# Helpers
def find_paragraph_index_containing(key):
    for i, p in enumerate(doc.paragraphs):
        if key in p.text:
            return i
    return None

# 1) Insert executive summary first sentence after the '1. Анотація' heading
idx = find_paragraph_index_containing('1. Анотація')
if idx is not None:
    # find next non-empty paragraph after heading
    j = idx + 1
    while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() == '':
        j += 1
    if j < len(doc.paragraphs):
        # Prepend the new first sentence to existing paragraph if not present
        if EXEC_SUM_FIRST not in doc.paragraphs[j].text:
            doc.paragraphs[j].text = EXEC_SUM_FIRST + '\n\n' + doc.paragraphs[j].text
            print('Updated Executive Summary at paragraph', j+1)
    else:
        # Append at end if no paragraph
        doc.add_paragraph(EXEC_SUM_FIRST)
        print('Appended Executive Summary at end')
else:
    print('Heading 1. Анотація not found')

# 2) Ensure Objective 4 contains analytics brief requirement
obj4_idx = find_paragraph_index_containing('Objective 4')
if obj4_idx is not None:
    if 'аналітич' not in doc.paragraphs[obj4_idx].text:
        doc.paragraphs[obj4_idx].text = doc.paragraphs[obj4_idx].text + ' ' + ANALYTICAL_BRIEFS
        print('Updated Objective 4 at paragraph', obj4_idx+1)
else:
    print('Objective 4 not found')

# 3) Insert budget note near a Budget heading if found, else append to end
budget_keys = ['Бюджет', 'Budget', '15.']
bidx = None
for k in budget_keys:
    bidx = find_paragraph_index_containing(k)
    if bidx is not None:
        break

if bidx is not None:
    # find next non-empty paragraph after budget heading
    j = bidx + 1
    while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() == '':
        j += 1
    # insert BUDGET_NOTE as a new paragraph before j
    p = doc.paragraphs[j-1]._element
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    new_p = doc.paragraphs[j-1]._element.addnext(p.__class__())
    new_par = doc.paragraphs[j-1]._parent.paragraphs[j]
    # fallback simple insertion
    doc.paragraphs[j-1]._parent.add_paragraph(BUDGET_NOTE)
    print('Inserted budget note after paragraph', j)
else:
    doc.add_paragraph(BUDGET_NOTE)
    print('Appended budget note at end')

# Save as new file
doc.save(OUTPUT)
print('Saved patched document to', OUTPUT)
