import os
import json
from pathlib import Path
import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Шляхи до файлів конфігурації та збереження
BASE_DIR = Path(__file__).resolve().parent
LEGAL_CONFIG_PATH = BASE_DIR / ".agents" / "skills" / "03-legal" / "legal_standards.json"
OUT_DIR = BASE_DIR / "_DROPZONE" / "OUT"
OUT_DIR.mkdir(parents=True, exist_ok=True)

def load_legal_standards():
    """Завантаження жорстких правил ДСТУ 4163:2020 та констант з JSON-конфігу"""
    if not LEGAL_CONFIG_PATH.exists():
        raise FileNotFoundError(f"❌ Конфігураційний файл не знайдено за адресою: {LEGAL_CONFIG_PATH}")
    with open(LEGAL_CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def create_styled_document(standards):
    """Створення документа з суворим застосуванням геометрії та типографіки ДСТУ"""
    doc = docx.Document()
    
    # 1. Налаштування полів (Margins)
    geom = standards["document_geometry"]["margins_mm"]
    for section in doc.sections:
        section.top_margin = Mm(geom["top"])
        section.bottom_margin = Mm(geom["bottom"])
        section.left_margin = Mm(geom["left"])
        section.right_margin = Mm(geom["right"])
        
    # 2. Налаштування базового шрифту та інтервалів (Typography)
    typo = standards["typography"]
    style = doc.styles['Normal']
    font = style.font
    font.name = typo["primary_font"]
    font.size = Pt(typo["base_font_size_pt"])
    font.color.rgb = RGBColor(*typo["text_color_rgb"])
    style.paragraph_format.line_spacing = typo["line_spacing"]
    style.paragraph_format.space_after = Pt(typo["paragraph_space_after_pt"])
    
    return doc

def add_table_row(table, data, is_header=False, widths=None):
    """Додавання рядка до таблиці з налаштуванням відступів, вирівнювання та жорсткої ширини"""
    row = table.add_row()
    for i, text in enumerate(data):
        cell = row.cells[i]
        cell.text = str(text)
        
        # Жорстко задаємо ширину для кожної клітинки кожного рядка (захист від warping у MS Word)
        if widths and i < len(widths):
            cell.width = widths[i]
            
        p = cell.paragraphs[0]
        # Робимо таблицю суперкомпактною, мінімізуючи вертикальні відступи всередині клітинок
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        
        # Налаштування стилю заголовка
        if is_header:
            for run in p.runs:
                run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Вирівнювання: № та кількісні дані по центру, найменування по лівому краю
            if i in [0, 2, 3, 4, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return row

def generate_transfer_act():
    try:
        # Завантажуємо стандарти
        standards = load_legal_standards()
        org_const = standards["organization_constants"]
        layout_rules = standards["layout_rules"]
        
        # Створюємо документ
        doc = create_styled_document(standards)
        
        # 1. Заголовок акту (компактні відступи)
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(6)
        run_title = p_title.add_run("АКТ ПРИЙМАННЯ-ПЕРЕДАЧІ БЛАГОДІЙНОЇ ДОПОМОГИ № 1/26")
        run_title.font.bold = True
        run_title.font.size = Pt(12)
        
        # 2. Дата та місце (компактні відступи)
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(12)
        run_date = p_meta.add_run("«___» ____________ 2026 року")
        p_meta.add_run("\t" * 6)  # Рознесення дати та місця в один рядок
        run_place = p_meta.add_run("м. Черкаси")
        run_date.font.italic = True
        run_place.font.italic = True
        
        # 3. Вступна преамбула (компактний абзац)
        p_preamble = doc.add_paragraph()
        p_preamble.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_preamble.paragraph_format.space_after = Pt(8)
        
        # Сторона 1 (ГО «ТАЛАН ЮА»)
        p_preamble.add_run("ГРОМАДСЬКА ОРГАНІЗАЦІЯ «ТАЛАН ЮА»").bold = True
        p_preamble.add_run(
            f" (код ЄДРПОУ: {org_const['edrpou']}, юридична адреса: {org_const['legal_address']}), "
            "в особі Голови правління "
        )
        p_preamble.add_run("Шаповала Тимура Юрійовича").bold = True
        p_preamble.add_run(
            ", що діє на підставі Статуту (надалі за текстом — «Передавач» або «Сторона 1»), з однієї сторони, та\n"
        )
        
        # Сторона 2 (Отримувач - ГО Партнер)
        p_preamble.add_run("[НАЗВА ГРОМАДСЬКОЇ ОРГАНІЗАЦІЯ - ОТРИМУВАЧА]").bold = True
        p_preamble.add_run(
            " (код ЄДРПОУ: [КОД ЄДРПОУ], юридична адреса: [ЮРИДИЧНА АДРЕСА]), "
            "в особі [Посада керівника Отримувача] "
        )
        p_preamble.add_run("[Прізвище, Ім'я, По батькові керівника Отримувача]").bold = True
        p_preamble.add_run(
            ", що діє на підставі [Статуту / Довіреності] (надалі за текстом — «Отримувач» або «Сторона 2»), з другої сторони, "
            "(в подальшому разом іменовані як «Сторони», а кожна окремо як «Сторона»), "
            "склали цей Акт приймання-передачі (надалі за текстом — «Акт») про наступне:"
        )
        
        # 4. Текст пунктів
        p_p1 = doc.add_paragraph()
        p_p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_p1.paragraph_format.space_after = Pt(6)
        p_p1.add_run("1. ").bold = True
        p_p1.add_run("Передавач передає, а Отримувач приймає у власність як благодійну (гуманітарну) допомогу наступне майно (товари):")
        
        # 5. Таблиця товарів (СУВОРІ ШИРИНИ ДЛЯ ВСІХ РЯДКІВ)
        # Сумарна ширина = 10 + 80 + 15 + 20 + 22 + 23 = 170 мм (ідеально під поля 30мм/10мм на А4)
        widths = [Mm(10), Mm(80), Mm(15), Mm(20), Mm(22), Mm(23)]
        
        table = doc.add_table(rows=0, cols=6, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Шапка таблиці
        add_table_row(table, ["№\nз/п", "Найменування майна (товару)", "Од.\nвим.", "Кіль-\nкість", "Вартість за\nод., грн", "Загальна\nвартість, грн"], is_header=True, widths=widths)
        
        # Створюємо ЧИСТІ порожні рядки для самостійного заповнення (як просив користувач!)
        add_table_row(table, ["1", "", "", "", "", ""], is_header=False, widths=widths)
        add_table_row(table, ["2", "", "", "", "", ""], is_header=False, widths=widths)
        add_table_row(table, ["3", "", "", "", "", ""], is_header=False, widths=widths)
        
        # Рядок "Разом" із жорстким об'єднанням та правильним вирівнюванням
        row_total = table.add_row()
        row_total.cells[0].merge(row_total.cells[4])
        row_total.cells[0].text = "ЗАГАЛЬНА ВАРТІСТЬ (БЕЗ ПДВ):"
        p_tot = row_total.cells[0].paragraphs[0]
        p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_tot.paragraph_format.space_before = Pt(2)
        p_tot.paragraph_format.space_after = Pt(2)
        p_tot.runs[0].font.bold = True
        
        row_total.cells[5].text = ""  # Порожнє поле для заповнення
        p_sum = row_total.cells[5].paragraphs[0]
        p_sum.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sum.paragraph_format.space_before = Pt(2)
        p_sum.paragraph_format.space_after = Pt(2)
        
        # Обов'язково встановлюємо ширину для об'єднаного рядка разом!
        row_total.cells[0].width = sum(widths[:5])
        row_total.cells[5].width = widths[5]
        
        doc.add_paragraph().paragraph_format.space_after = Pt(2)  # Мінімальний відступ після таблиці
        
        # 6. Наступні пункти акту (оптимізовані під 1 сторінку)
        p_p2 = doc.add_paragraph()
        p_p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_p2.paragraph_format.space_after = Pt(4)
        p_p2.add_run("2. ").bold = True
        p_p2.add_run("Загальна вартість майна, що передається за цим Актом, становить ")
        p_p2.add_run("_________________________________________ грн. (______________________________________________________) ").bold = True
        p_p2.add_run("без ПДВ.")
        
        p_p3 = doc.add_paragraph()
        p_p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_p3.paragraph_format.space_after = Pt(4)
        p_p3.add_run("3. ").bold = True
        p_p3.add_run(
            "Благодійна допомога передається Отримувачу для використання виключно в статутній некомерційній діяльності, "
            "спрямованій на досягнення цілей, передбачених статутом Отримувача та Законом України «Про благодійну діяльність та благодійні організації», "
            "і не може бути використана з метою отримання прибутку."
        )
        
        # Впроваджуємо жорстке правило утримання блоку підписів на одній сторінці (keep_with_next)
        p_p4 = doc.add_paragraph()
        p_p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_p4.paragraph_format.space_after = Pt(4)
        p_p4.paragraph_format.keep_with_next = True  # Обов'язкове склеювання сторінки
        p_p4.add_run("4. ").bold = True
        p_p4.add_run(
            "Отримувач підтверджує, що майно передано у належному стані, комплектне, видимих дефектів та пошкоджень не виявлено. "
            "Сторони підтверджують факт повного передавання майна та заявляють, що не мають одна до одної жодних претензій."
        )
        
        p_p5 = doc.add_paragraph()
        p_p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_p5.paragraph_format.space_after = Pt(12)
        p_p5.paragraph_format.keep_with_next = True  # Прив'язуємо безпосередньо до заголовка підписів!
        p_p5.add_run("5. ").bold = True
        p_p5.add_run(
            "Цей Акт складений українською мовою у двох оригінальних примірниках, які мають однакову юридичну силу, "
            "по одному примірнику для кожної із Сторін."
        )
        
        # 7. Блок підписів за ДСТУ
        sig_styling = layout_rules["signature_block_styling"]
        
        p_sig_head = doc.add_paragraph()
        p_sig_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sig_head.paragraph_format.space_before = Pt(8)
        p_sig_head.paragraph_format.space_after = Pt(8)
        p_sig_head.paragraph_format.keep_with_next = True
        
        run_sig_head = p_sig_head.add_run(sig_styling["header_text"])
        run_sig_head.font.bold = sig_styling["header_bold"]
        
        # Створення таблиці підписів (2 стовпці) — максимально компактний розмір
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.autofit = False
        sig_table.columns[0].width = Mm(85)
        sig_table.columns[1].width = Mm(85)
        
        # Рядок 1: Назва та Посади (Сторона 1 - Ліво, Сторона 2 - Право)
        cell_l0 = sig_table.cell(0, 0)
        cell_l0.text = "ПЕРЕДАВАЧ:\nГолова правління\nГО «ТАЛАН ЮА»"
        cell_l0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell_l0.paragraphs[0].runs[0].font.bold = True
        cell_l0.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell_l0.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell_l0.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        cell_r0 = sig_table.cell(0, 1)
        cell_r0.text = "ОТРИМУВАЧ:\n[Посада керівника]\n[НАЗВА ГО-ОТРИМУВАЧА]"
        cell_r0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell_r0.paragraphs[0].runs[0].font.bold = True
        cell_r0.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell_r0.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell_r0.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        # Рядок 2: Місце підпису та Ім'я/Прізвище (Прізвище ВЕЛИКИМИ ЛІТЕРАМИ за ДСТУ, без зайвих порожніх ліній)
        cell_l1 = sig_table.cell(1, 0)
        cell_l1.text = "\n_________________ / Тимур ШАПОВАЛ /\nМ.П."
        cell_l1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell_l1.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell_l1.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell_l1.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        cell_r1 = sig_table.cell(1, 1)
        cell_r1.text = "\n_________________ / [Ініціали ПРІЗВИЩЕ] /\nМ.П."
        cell_r1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell_r1.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell_r1.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell_r1.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        # Спроба збереження з авто-інкрементом версії у разі блокування (PermissionError)
        output_filename = "08_Act_of_Transfer_TalanUA_TEMPLATE.docx"
        output_path = OUT_DIR / output_filename
        
        version = 1
        while True:
            try:
                doc.save(output_path)
                break
            except PermissionError:
                version += 1
                output_filename = f"08_Act_of_Transfer_TalanUA_TEMPLATE_v{version}.docx"
                output_path = OUT_DIR / output_filename
                
        print(f"🎉 Акт прийому-передачі успішно згенеровано: {output_path}")
        
    except Exception as e:
        print(f"❌ Помилка під час генерації акту: {e}")

if __name__ == "__main__":
    generate_transfer_act()
